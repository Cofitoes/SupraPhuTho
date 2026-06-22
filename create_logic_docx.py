from docx import Document
from docx.shared import Pt
import os

md_path = r"g:\My Drive\Training AI\Supra Phú Thọ\Logic_Chia_Tuyen_Du_An.md"
docx_path = r"g:\My Drive\Training AI\Supra Phú Thọ\Logic_Chia_Tuyen.docx"

doc = Document()

# Add a Title
title = doc.add_heading('TÀI LIỆU LOGIC CHIA TUYẾN VẬN TẢI', 0)

with open(md_path, 'r', encoding='utf-8') as f:
    content = f.read()

lines = content.split('\n')
for line in lines:
    if line.startswith('# '):
        continue  # skip title
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=1)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=2)
    elif line.startswith('**') and line.endswith('**'):
        p = doc.add_paragraph()
        p.add_run(line[2:-2]).bold = True
    elif line.startswith('* **') or line.startswith('- **'):
        p = doc.add_paragraph(style='List Bullet')
        parts = line[4:].split('**', 1)
        if len(parts) == 2:
            p.add_run(parts[0]).bold = True
            p.add_run(parts[1])
        else:
            p.add_run(line)
    elif line.startswith('* ') or line.startswith('- '):
        doc.add_paragraph(line[2:], style='List Bullet')
    elif line.strip() == '---':
        continue
    else:
        if line.strip() != "":
            doc.add_paragraph(line)

doc.save(docx_path)
print("Created docx file successfully!")
