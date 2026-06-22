from docx import Document
from docx.shared import Pt

docx_path = r'g:\My Drive\Training AI\Supra Phú Thọ\Logic_Chia_Tuyen.docx'
doc = Document(docx_path)

# Clear everything after paragraph 10 just to be safe, though it's already empty
p_count = len(doc.paragraphs)
if p_count > 11:
    for i in range(p_count - 1, 10, -1):
        p = doc.paragraphs[i]
        p._element.getparent().remove(p._element)

# Add the new logic for Section 2
doc.add_paragraph('Quy trình tối ưu hóa và ghép chuyến được thực hiện theo nguyên tắc sau:', style='Normal')

p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('Điều phối về kho GXT Phú Thọ: ').bold = True
p1.add_run('Dựa vào cột I trong danh sách (sheet Data), những siêu thị nào được đánh dấu ')
p1.add_run('"GXT"').bold = True
p1.add_run(' sẽ được tự động gom và vận chuyển về kho GXT Phú Thọ. Hệ thống sẽ tự động tính toán tổng lượng hàng, trọng tải và thể tích của nhóm siêu thị này để đề xuất loại xe (1.9T, 5T, 8T) phù hợp nhất, đảm bảo tối ưu chi phí.')

p2 = doc.add_paragraph(style='List Bullet')
p2.add_run('Giao thẳng trực tiếp từ Hub chính (Kho DC Win Phúc Thọ): ').bold = True
p2.add_run('Đối với tất cả các siêu thị khác (không nằm trong danh sách GXT), hàng sẽ được giao trực tiếp từ Hub chính. Thuật toán chia tuyến sẽ dựa trên 3 tiêu chí cốt lõi:\n')
p2.add_run('   • Trọng tải (kg)\n')
p2.add_run('   • Thể tích (CBM)\n')
p2.add_run('   • Tọa độ vị trí (Location) của cửa hàng (ưu tiên ghép các điểm giao nằm trên cùng một tuyến đường hoặc có vị trí địa lý gần nhau).')

doc.save(docx_path)
print("Updated DOCX successfully!")
