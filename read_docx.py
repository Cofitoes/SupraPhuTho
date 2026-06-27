import docx
import os

doc_path = r"g:\My Drive\Training AI\Supra Phú Thọ\Logic Chia Tuyen New.docx"
if os.path.exists(doc_path):
    doc = docx.Document(doc_path)
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"[{i}]: {para.text}")
            
    # Also check tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if "ConCung" in cell.text or "Con Cưng" in cell.text:
                    print(f"[Table {t_idx} Row {r_idx} Col {c_idx}]: {cell.text}")
else:
    print("DOCX file not found.")
